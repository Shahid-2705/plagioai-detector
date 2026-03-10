"""
Report Generator Module
Creates PDF and DOCX plagiarism reports.
"""

import os
import datetime
import tempfile


def generate_report(
    rewritten_sentences: list,
    original_score: float,
    new_score: float,
    format_type: str = 'pdf',
) -> str:
    """
    Generate a plagiarism report in the requested format.

    Parameters
    ----------
    rewritten_sentences : list of dicts with keys:
        original, rewritten, was_rewritten, original_score
    original_score : float (0-100)
    new_score : float (0-100)
    format_type : 'pdf' | 'docx'

    Returns
    -------
    str : filepath of the generated report
    """
    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    if format_type == 'pdf':
        return _generate_pdf(rewritten_sentences, original_score, new_score, timestamp)
    elif format_type == 'docx':
        return _generate_docx(rewritten_sentences, original_score, new_score, timestamp)
    else:
        raise ValueError(f"Unsupported format: {format_type}")


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _generate_pdf(sentences, original_score, new_score, timestamp):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        )

        filepath = os.path.join(tempfile.gettempdir(), f'plagiarism_report_{timestamp}.pdf')
        doc = SimpleDocTemplate(filepath, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle('Title', parent=styles['Title'],
                                     fontSize=20, textColor=colors.HexColor('#1a1a2e'),
                                     spaceAfter=6)
        subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'],
                                        fontSize=11, textColor=colors.HexColor('#555'),
                                        spaceAfter=12)
        heading_style = ParagraphStyle('Heading', parent=styles['Heading2'],
                                       fontSize=13, textColor=colors.HexColor('#1a1a2e'),
                                       spaceBefore=12, spaceAfter=6)
        body_style = ParagraphStyle('Body', parent=styles['Normal'],
                                    fontSize=10, leading=14)
        high_style = ParagraphStyle('High', parent=body_style,
                                    textColor=colors.HexColor('#c0392b'))
        rewritten_style = ParagraphStyle('Rewritten', parent=body_style,
                                         textColor=colors.HexColor('#27ae60'))

        story = []

        # Title
        story.append(Paragraph('AI Plagiarism Detector & Rewriter', title_style))
        story.append(Paragraph(f'Report generated: {datetime.datetime.now().strftime("%B %d, %Y %H:%M")}', subtitle_style))
        story.append(HRFlowable(width='100%', thickness=1, color=colors.HexColor('#ddd')))
        story.append(Spacer(1, 12))

        # Summary table
        story.append(Paragraph('Summary', heading_style))
        reduction = original_score - new_score
        summary_data = [
            ['Metric', 'Value'],
            ['Original Plagiarism Score', f'{original_score:.1f}%'],
            ['New Plagiarism Score', f'{new_score:.1f}%'],
            ['Score Reduction', f'{reduction:.1f}%'],
            ['Sentences Rewritten', str(sum(1 for s in sentences if s.get('was_rewritten')))],
            ['Total Sentences', str(len(sentences))],
        ]
        t = Table(summary_data, colWidths=[10*cm, 7*cm])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1a1a2e')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 10),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor('#f8f8f8'), colors.white]),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#ddd')),
            ('PADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 16))

        # Sentence-by-sentence analysis
        story.append(Paragraph('Sentence Analysis', heading_style))
        story.append(HRFlowable(width='100%', thickness=0.5, color=colors.HexColor('#eee')))
        story.append(Spacer(1, 6))

        for i, s in enumerate(sentences):
            score_pct = s.get('original_score', 0) * 100
            label = 'HIGH RISK' if score_pct >= 70 else ('MEDIUM RISK' if score_pct >= 40 else 'LOW RISK')
            label_color = '#c0392b' if score_pct >= 70 else ('#e67e22' if score_pct >= 40 else '#27ae60')

            story.append(Paragraph(
                f'<b>Sentence {i+1}</b>  '
                f'<font color="{label_color}">● {label} ({score_pct:.0f}%)</font>',
                body_style
            ))
            story.append(Paragraph(f'<b>Original:</b> {s["original"]}', high_style if score_pct >= 70 else body_style))
            if s.get('was_rewritten'):
                story.append(Paragraph(f'<b>Rewritten:</b> {s["rewritten"]}', rewritten_style))
            story.append(Spacer(1, 8))

        doc.build(story)
        return filepath

    except ImportError:
        raise ImportError("reportlab is required. Run: pip install reportlab")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _generate_docx(sentences, original_score, new_score, timestamp):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        filepath = os.path.join(tempfile.gettempdir(), f'plagiarism_report_{timestamp}.docx')
        doc = Document()

        # Title
        title = doc.add_heading('AI Plagiarism Detector & Rewriter', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph(f'Report generated: {datetime.datetime.now().strftime("%B %d, %Y %H:%M")}')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('Summary', level=1)
        reduction = original_score - new_score
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        headers = ['Metric', 'Value']
        for j, h in enumerate(headers):
            table.cell(0, j).text = h
        rows_data = [
            ('Original Plagiarism Score', f'{original_score:.1f}%'),
            ('New Plagiarism Score', f'{new_score:.1f}%'),
            ('Score Reduction', f'{reduction:.1f}%'),
            ('Sentences Rewritten', str(sum(1 for s in sentences if s.get('was_rewritten')))),
            ('Total Sentences', str(len(sentences))),
        ]
        for row_idx, (k, v) in enumerate(rows_data, start=1):
            table.cell(row_idx, 0).text = k
            table.cell(row_idx, 1).text = v

        doc.add_heading('Sentence Analysis', level=1)
        for i, s in enumerate(sentences):
            score_pct = s.get('original_score', 0) * 100
            label = 'HIGH RISK' if score_pct >= 70 else ('MEDIUM RISK' if score_pct >= 40 else 'LOW RISK')
            p = doc.add_paragraph()
            p.add_run(f'Sentence {i+1} — ').bold = True
            run = p.add_run(f'{label} ({score_pct:.0f}%)')
            run.bold = True
            if score_pct >= 70:
                run.font.color.rgb = RGBColor(192, 57, 43)
            elif score_pct >= 40:
                run.font.color.rgb = RGBColor(230, 126, 34)
            else:
                run.font.color.rgb = RGBColor(39, 174, 96)

            orig_p = doc.add_paragraph()
            orig_p.add_run('Original: ').bold = True
            orig_p.add_run(s['original'])

            if s.get('was_rewritten'):
                rew_p = doc.add_paragraph()
                rew_p.add_run('Rewritten: ').bold = True
                rew_run = rew_p.add_run(s['rewritten'])
                rew_run.font.color.rgb = RGBColor(39, 174, 96)

        doc.save(filepath)
        return filepath

    except ImportError:
        raise ImportError("python-docx is required. Run: pip install python-docx")
