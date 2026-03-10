"""
File Parser Module
Extracts text from PDF, DOCX, and TXT files.
"""

import os


def extract_text_from_file(filepath: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        return extract_pdf_text(filepath)
    elif ext == '.docx':
        return extract_docx_text(filepath)
    elif ext == '.txt':
        return extract_txt_text(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def extract_pdf_text(filepath: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text.strip())
        return '\n\n'.join(text_parts)
    except ImportError:
        raise ImportError("pdfplumber is required. Run: pip install pdfplumber")
    except Exception as e:
        raise RuntimeError(f"Failed to extract PDF text: {e}")


def extract_docx_text(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return '\n\n'.join(paragraphs)
    except ImportError:
        raise ImportError("python-docx is required. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Failed to extract DOCX text: {e}")


def extract_txt_text(filepath: str) -> str:
    """Extract text from a plain text file."""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
    for encoding in encodings:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise RuntimeError("Could not decode text file with any supported encoding.")
