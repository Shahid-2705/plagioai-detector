"""
PlagioAI — app.py  (Production build)

Deployment:
  Frontend → Vercel
  Backend  → Render  (https://plagioai-detector.onrender.com)
  Session  → in-memory dict (no filesystem dependency on Render free tier)

Design decisions:
  - Models loaded ONCE at module import in background threads (never per-request)
  - CORS allows any origin so Vercel frontend can call Render backend
  - Session replaced with a simple in-memory dict keyed by session_id cookie
    (avoids flask-session filesystem writes on read-only Render containers)
  - /status returns granular sub-status so the banner shows download progress
  - Adaptive thresholding: < 8 words → 0.40, 8-20 → 0.55, > 20 → 0.70
  - /download_rewritten returns full merged doc as DOCX or TXT
"""

import os
import re
import uuid
import time
import tempfile
import warnings
import threading
import traceback
import logging

warnings.filterwarnings('ignore', message='.*sequentially on GPU.*', category=UserWarning)
logging.basicConfig(level=logging.INFO)

from flask import Flask, request, jsonify, send_file, make_response
from flask_cors import CORS

from core.file_parser         import extract_text_from_file
from core.text_preprocessor   import segment_sentences
from core.plagiarism_detector  import PlagiarismDetector
from core.rewrite_engine       import RewriteEngine
from core.report_generator     import generate_report

# ── Detect GPU ────────────────────────────────────────────────────────────────
try:
    import torch
    _GPU = torch.cuda.get_device_name(0) if torch.cuda.is_available() else 'CPU'
    print(f'[PlagioAI] Device: {_GPU}')
except Exception:
    _GPU = 'CPU'
    print('[PlagioAI] Device: CPU')

# ── Flask + CORS ───────────────────────────────────────────────────────────────
app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', os.urandom(24))

# Allow any origin (Vercel preview URLs change per deploy)
CORS(app,
     supports_credentials=True,
     origins='*',
     allow_headers=['Content-Type', 'Authorization'],
     methods=['GET', 'POST', 'OPTIONS'])

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

# ── In-memory session store ────────────────────────────────────────────────────
# Keyed by a UUID stored in the 'sid' cookie.
# Each entry: { 'original_text': str, 'sentences': list,
#               'detection_results': dict, 'rewritten_sentences': list,
#               'new_score': float, 'ts': float }
_SESSIONS: dict = {}
_SESSION_TTL    = 3600   # seconds before a session is evicted

def _gc_sessions():
    """Remove sessions older than TTL."""
    cutoff = time.time() - _SESSION_TTL
    dead   = [k for k, v in _SESSIONS.items() if v.get('ts', 0) < cutoff]
    for k in dead:
        _SESSIONS.pop(k, None)

def _get_sid():
    """Return existing sid from cookie, or create a new one."""
    return request.cookies.get('sid')

def _get_session() -> dict:
    sid = _get_sid()
    if sid and sid in _SESSIONS:
        _SESSIONS[sid]['ts'] = time.time()
        return _SESSIONS[sid]
    return {}

def _save_session(data: dict):
    """Persist data into the session, creating it if needed."""
    sid = _get_sid()
    if not sid or sid not in _SESSIONS:
        sid = str(uuid.uuid4())
    _SESSIONS.setdefault(sid, {})
    _SESSIONS[sid].update(data)
    _SESSIONS[sid]['ts'] = time.time()
    _gc_sessions()
    return sid

def _set_sid_cookie(response, sid: str):
    response.set_cookie(
        'sid', sid,
        max_age=_SESSION_TTL,
        samesite='None',
        secure=True,      # required for cross-site cookies on HTTPS
        httponly=True,
    )
    return response

# ── Skip / threshold helpers ───────────────────────────────────────────────────
_SKIP_PAT = re.compile(
    r'\b(figure|fig\.|table|tbl\.|references?|bibliography|doi|et\s+al\.?'
    r'|journal|proc\.|proceedings|vol\.|pp\.|isbn|issn|rrn|dept|course\s+code'
    r'|http|https|www\.)\b|10\.\d{4,}|https?://',
    re.IGNORECASE,
)
_NUMS_ONLY = re.compile(r'^[\d\s\.\,\:\;\-\+\%\(\)\/\=\u00b0\u03bc\u00b1\u00d7\u00f7]+$')

def _is_non_rewritable(s: str) -> bool:
    s = s.strip()
    return len(s) < 20 or bool(_SKIP_PAT.search(s)) or bool(_NUMS_ONLY.match(s))

def _adaptive_threshold(s: str) -> float:
    wc = len(s.split())
    if wc < 8:   return 0.40
    if wc <= 20: return 0.55
    return 0.70

# ── Model registry ─────────────────────────────────────────────────────────────
_models: dict = {
    'detector':        None,
    'rewriter':        None,
    'detector_ready':  False,
    'rewriter_ready':  False,
    'detector_status': 'loading',
    'rewriter_status': 'loading',
    'detector_sub':    'Starting...',
    'rewriter_sub':    'Starting...',
    'detector_error':  '',
    'rewriter_error':  '',
}
_lock = threading.Lock()

def _set(key, val):
    with _lock:
        _models[key] = val

# ── Background loader: detector ────────────────────────────────────────────────
def _load_detector():
    try:
        _set('detector_sub', 'Importing libraries...')
        _set('detector_sub', 'Loading embedding model (first run ~90-420 MB)...')
        d = PlagiarismDetector()
        _set('detector_sub', 'Warming up...')
        d.engine.compute_embeddings(['warmup'])
        model_name = getattr(d.engine, '_loaded_model_name', 'embedding model')
        with _lock:
            _models['detector']        = d
            _models['detector_ready']  = True
            _models['detector_status'] = 'ready'
            _models['detector_sub']    = f'Ready ({model_name})'
        print(f'[PlagioAI] Detector ready — {model_name}')
    except Exception:
        err = traceback.format_exc()
        short = err.strip().splitlines()[-1]
        with _lock:
            _models['detector_status'] = 'error'
            _models['detector_sub']    = short
            _models['detector_error']  = err
        print(f'[PlagioAI] Detector FAILED:\n{err}')

# ── Background loader: rewriter ────────────────────────────────────────────────
def _load_rewriter():
    try:
        _set('rewriter_sub', 'Importing libraries...')
        _set('rewriter_sub', 'Loading flan-t5-large (first run ~1 GB)...')
        r = RewriteEngine()
        _set('rewriter_sub', 'Warming up rewriter...')
        r.rewrite('Warm-up sentence.')
        with _lock:
            _models['rewriter']        = r
            _models['rewriter_ready']  = True
            _models['rewriter_status'] = 'ready'
            _models['rewriter_sub']    = 'Ready (flan-t5-large)'
        print('[PlagioAI] Rewriter ready')
    except Exception:
        err = traceback.format_exc()
        short = err.strip().splitlines()[-1]
        with _lock:
            _models['rewriter_status'] = 'error'
            _models['rewriter_sub']    = short
            _models['rewriter_error']  = err
        print(f'[PlagioAI] Rewriter FAILED:\n{err}')

print('[PlagioAI] Spawning model-load threads...')
threading.Thread(target=_load_detector, daemon=True).start()
threading.Thread(target=_load_rewriter, daemon=True).start()

def get_detector(): return _models.get('detector')
def get_rewriter():  return _models.get('rewriter')

def allowed_file(fn: str) -> bool:
    return '.' in fn and fn.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# ── Document helpers ───────────────────────────────────────────────────────────
def _build_rewritten_text(rows: list) -> str:
    return ' '.join(
        (r.get('rewritten') or r.get('original') or '').strip()
        for r in rows
        if (r.get('rewritten') or r.get('original') or '').strip()
    )

def _save_rewritten_docx(text: str, path: str):
    from docx import Document
    doc = Document()
    doc.add_heading('Rewritten Document', 0)
    doc.add_paragraph('Generated by PlagioAI')
    doc.add_paragraph('')
    for p in (text.split('\n\n') if '\n\n' in text else [text]):
        if p.strip():
            doc.add_paragraph(p.strip())
    doc.save(path)

# ── Health check / status ──────────────────────────────────────────────────────

@app.route('/health')
@app.route('/')
def health():
    """Render pings / for health checks."""
    return jsonify({'status': 'running', 'gpu': _GPU})


@app.route('/status')
def model_status():
    """
    Polled every 2 s by the frontend.
    Returns per-model status + sub-status message for live banner updates.
    """
    with _lock:
        ds  = _models['detector_status']
        rs  = _models['rewriter_status']
        return jsonify({
            'status':          'running',
            'detector_status': ds,
            'rewriter_status': rs,
            'detector_sub':    _models['detector_sub'],
            'rewriter_sub':    _models['rewriter_sub'],
            'detector_error':  _models['detector_error'],
            'rewriter_error':  _models['rewriter_error'],
            'all_ready':       ds == 'ready' and rs == 'ready',
            'ready':           ds == 'ready' and rs == 'ready',
        })


@app.route('/debug')
def debug_info():
    """Open /debug in browser to see exact error details."""
    with _lock:
        lines = [
            f"detector_status : {_models['detector_status']}",
            f"detector_sub    : {_models['detector_sub']}",
            f"rewriter_status : {_models['rewriter_status']}",
            f"rewriter_sub    : {_models['rewriter_sub']}",
            f"device          : {_GPU}",
        ]
        de  = _models['detector_error']
        re_ = _models['rewriter_error']

    for pkg in ('sentence_transformers', 'transformers', 'torch'):
        try:
            m = __import__(pkg)
            lines.append(f'{pkg:30s}: {m.__version__}')
        except Exception as e:
            lines.append(f'{pkg:30s}: NOT FOUND ({e})')

    lines += ['', '=== Detector error ===', de or '(none)',
              '', '=== Rewriter error ===', re_ or '(none)']
    return '<pre style="font:14px monospace;padding:2rem;background:#111;color:#eee;">' \
           + '\n'.join(lines) + '</pre>'


# ── Upload ─────────────────────────────────────────────────────────────────────

@app.route('/upload', methods=['POST', 'OPTIONS'])
def upload_file():
    if request.method == 'OPTIONS':
        return _cors_preflight()

    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400
    if not allowed_file(file.filename):
        return jsonify({'error': 'Unsupported file type. Use PDF, DOCX, or TXT.'}), 400

    from werkzeug.utils import secure_filename
    fn   = secure_filename(file.filename)
    path = os.path.join(UPLOAD_DIR, f'{uuid.uuid4()}_{fn}')
    file.save(path)

    try:
        text = extract_text_from_file(path)
        if not text or len(text.strip()) < 50:
            return jsonify({'error': 'Could not extract enough text.'}), 400

        sid = _save_session({'original_text': text})
        resp = make_response(jsonify({
            'success':    True,
            'full_text':  text,
            'preview':    text[:500] + '...' if len(text) > 500 else text,
            'word_count': len(text.split()),
            'char_count': len(text),
            'filename':   fn,
        }))
        return _set_sid_cookie(resp, sid)

    except Exception as e:
        return jsonify({'error': f'Error processing file: {e}'}), 500


# ── Detect ─────────────────────────────────────────────────────────────────────

@app.route('/detect', methods=['POST', 'OPTIONS'])
def detect_plagiarism():
    if request.method == 'OPTIONS':
        return _cors_preflight()

    detector = get_detector()
    if detector is None:
        with _lock:
            st  = _models['detector_status']
            sub = _models['detector_sub']
            err = _models['detector_error']
        if st == 'error':
            return jsonify({'error': f'Detector failed: {err.strip().splitlines()[-1]}'}), 503
        return jsonify({'error': f'Detector loading: {sub}'}), 503

    body    = request.get_json(silent=True) or {}
    sess    = _get_session()
    text    = (body.get('text') or '').strip() or (sess.get('original_text') or '').strip()
    if not text:
        return jsonify({'error': 'No text to analyse. Upload or paste text first.'}), 400

    try:
        sentences = segment_sentences(text)
        if len(sentences) < 2:
            return jsonify({'error': 'Need at least 2 sentences.'}), 400

        results = detector.analyze(sentences)
        sid     = _save_session({'sentences': sentences, 'detection_results': results})
        resp    = make_response(jsonify({
            'success':           True,
            'sentences':         sentences,
            'results':           results['sentence_scores'],
            'overall_score':     results['overall_score'],
            'high_risk_count':   results['high_risk_count'],
            'medium_risk_count': results['medium_risk_count'],
        }))
        return _set_sid_cookie(resp, sid)

    except Exception as e:
        return jsonify({'error': f'Detection error: {e}'}), 500


# ── Rewrite ────────────────────────────────────────────────────────────────────

@app.route('/rewrite', methods=['POST', 'OPTIONS'])
def rewrite_content():
    if request.method == 'OPTIONS':
        return _cors_preflight()

    rewriter = get_rewriter()
    if rewriter is None:
        with _lock:
            st  = _models['rewriter_status']
            sub = _models['rewriter_sub']
            err = _models['rewriter_error']
        if st == 'error':
            return jsonify({'error': f'Rewriter failed: {err.strip().splitlines()[-1]}'}), 503
        return jsonify({'error': f'Rewriter loading: {sub}'}), 503

    body = request.get_json(silent=True) or {}
    sess = _get_session()

    sentences   = body.get('sentences')   or sess.get('sentences', [])
    stored      = sess.get('detection_results', {})
    sent_scores = stored.get('sentence_scores', [])
    overall_in  = float(stored.get('overall_score', 0))

    if not sent_scores:
        raw         = body.get('results') or {}
        sent_scores = raw.get('sentence_scores', [])
        overall_in  = float(raw.get('overall_score', overall_in))

    if not sentences:
        return jsonify({'error': 'No sentences found. Run detection first.'}), 400

    try:
        scores = [
            float((sent_scores[i] if i < len(sent_scores) else {}).get('combined_score', 0))
            for i in range(len(sentences))
        ]

        flagged_idx, flagged_sents, thresholds = [], [], []
        for i, s in enumerate(sentences):
            thr = _adaptive_threshold(s)
            thresholds.append(thr)
            if scores[i] >= thr and not _is_non_rewritable(s):
                flagged_idx.append(i)
                flagged_sents.append(s)

        batch_out   = rewriter.rewrite_batch(flagged_sents) if flagged_sents else []
        rewrite_map = dict(zip(flagged_idx, batch_out))

        rows = []
        for i, s in enumerate(sentences):
            rt = rewrite_map.get(i, s)
            rw = (i in rewrite_map
                  and rt.strip().lower() != s.strip().lower()
                  and len(rt.strip()) > 10)
            rows.append({
                'original':       s,
                'rewritten':      rt if rw else s,
                'was_rewritten':  rw,
                'original_score': scores[i],
                'threshold_used': thresholds[i],
                'non_rewritable': _is_non_rewritable(s),
            })

        new_text    = _build_rewritten_text(rows)
        new_sents   = segment_sentences(new_text)
        det         = get_detector()
        new_res     = det.analyze(new_sents) if det and len(new_sents) >= 2 else {'overall_score': 0}
        new_score   = new_res.get('overall_score', 0)

        sid  = _save_session({'rewritten_sentences': rows, 'new_score': new_score})
        resp = make_response(jsonify({
            'success':             True,
            'rewritten_sentences': rows,
            'original_score':      overall_in,
            'new_score':           new_score,
            'rewritten_count':     sum(1 for r in rows if r['was_rewritten']),
            'skipped_count':       sum(1 for r in rows if r['non_rewritable'] and r['original_score'] >= r['threshold_used']),
            'total_sentences':     len(sentences),
            'score_reduction':     round(overall_in - new_score, 2),
        }))
        return _set_sid_cookie(resp, sid)

    except Exception as e:
        return jsonify({'error': f'Rewriting error: {e}'}), 500


# ── Download rewritten document ────────────────────────────────────────────────

@app.route('/download_rewritten', methods=['POST', 'OPTIONS'])
def download_rewritten():
    """
    POST /download_rewritten
    Body: { "format": "docx" | "txt" }
    Returns the full merged rewritten document as a downloadable file.
    """
    if request.method == 'OPTIONS':
        return _cors_preflight()

    body = request.get_json(silent=True) or {}
    fmt  = body.get('format', 'docx').lower()
    sess = _get_session()
    rows = sess.get('rewritten_sentences', [])

    if not rows:
        return jsonify({'error': 'No rewritten document. Complete the rewrite step first.'}), 400

    full = _build_rewritten_text(rows)
    tid  = uuid.uuid4()

    try:
        if fmt == 'docx':
            fp   = os.path.join(tempfile.gettempdir(), f'rw_{tid}.docx')
            _save_rewritten_docx(full, fp)
            name = 'rewritten_document.docx'
            mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            fp   = os.path.join(tempfile.gettempdir(), f'rw_{tid}.txt')
            with open(fp, 'w', encoding='utf-8') as f:
                f.write('REWRITTEN DOCUMENT\nGenerated by PlagioAI\n' + '=' * 60 + '\n\n' + full)
            name = 'rewritten_document.txt'
            mime = 'text/plain'

        return send_file(fp, as_attachment=True, download_name=name, mimetype=mime)

    except Exception as e:
        return jsonify({'error': f'Download error: {e}'}), 500


# ── Report ─────────────────────────────────────────────────────────────────────

@app.route('/report', methods=['POST', 'OPTIONS'])
def download_report():
    if request.method == 'OPTIONS':
        return _cors_preflight()

    body        = request.get_json(silent=True) or {}
    format_type = body.get('format', 'pdf')
    orig_score  = float(body.get('original_score', 0))
    sess        = _get_session()
    rows        = sess.get('rewritten_sentences', [])
    new_score   = float(sess.get('new_score', 0))

    if not rows:
        return jsonify({'error': 'No data. Complete the rewrite step first.'}), 400

    try:
        fp = generate_report(rows, orig_score, new_score, format_type)
        return send_file(fp, as_attachment=True,
                         download_name=f'plagiarism_report.{format_type}')
    except Exception as e:
        return jsonify({'error': f'Report error: {e}'}), 500


# ── CORS preflight helper ──────────────────────────────────────────────────────

def _cors_preflight():
    resp = make_response('', 204)
    resp.headers['Access-Control-Allow-Origin']      = request.headers.get('Origin', '*')
    resp.headers['Access-Control-Allow-Credentials'] = 'true'
    resp.headers['Access-Control-Allow-Headers']     = 'Content-Type, Authorization'
    resp.headers['Access-Control-Allow-Methods']     = 'GET, POST, OPTIONS'
    return resp


# ── Entry point ────────────────────────────────────────────────────────────────



if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))

    print(f"Starting server on port {port}")

    app.run(
        host="0.0.0.0",
        port=port,
        debug=False
    )